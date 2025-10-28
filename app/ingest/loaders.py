import os, glob, re, hashlib, html2text, markdown, requests
from typing import List, Iterable
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader
from langchain_core.documents import Document

# Optional Readability extraction for articles (auto-disabled if missing)
try:
    from readability import Document as ReadabilityDoc
    _HAS_READABILITY = True
except Exception:
    _HAS_READABILITY = False

# Supported file extensions
TEXT_EXTS = {".txt", ".md", ".markdown", ".html", ".htm"}
DOCX_EXTS = {".docx"}
PDF_EXTS  = {".pdf"}

# Cleaning helpers
AD_STOPWORDS = {
    "subscribe", "sign up", "cookie", "privacy", "terms",
    "share this", "available at:", "myplate.gov", "myplate plan",
    "back to top", "newsletter", "sponsored"
}
URL_RE = re.compile(r"https?://\S+")
MULTI_WS_RE = re.compile(r"\s+")

def _norm_ws(s: str) -> str:
    return MULTI_WS_RE.sub(" ", s).strip()

def _looks_like_boilerplate(line: str) -> bool:
    low = line.lower()
    if len(low) < 30:
        return True
    if URL_RE.search(low):
        return True
    if any(tok in low for tok in AD_STOPWORDS):
        return True
    alpha = sum(c.isalpha() for c in low)
    return alpha < 10

def _strip_boilerplate(text: str) -> str:
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    keep = [l for l in lines if not _looks_like_boilerplate(l)]
    return _norm_ws("\n".join(keep))

def _hash_content(s: str) -> str:
    return hashlib.sha1(s.encode("utf-8", "ignore")).hexdigest()

# --- Loaders ---

def load_txt_like(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        raw = f.read()

    if ext in {".md", ".markdown"}:
        html = markdown.markdown(raw)
        text = html2text.html2text(html)
        return _strip_boilerplate(text)

    if ext in {".html", ".htm"}:
        soup = BeautifulSoup(raw, "html.parser")
        for tag in soup(["script", "style", "noscript", "header", "footer", "nav", "aside"]):
            tag.decompose()
        text = soup.get_text(separator=" ")
        return _strip_boilerplate(text)

    return _strip_boilerplate(raw)

def load_docx(path: str) -> str:
    doc = DocxDocument(path)
    parts = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    return _strip_boilerplate("\n".join(parts))

def load_pdf(path: str) -> Iterable[Document]:
    reader = PdfReader(path)
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        text = _strip_boilerplate(text)
        if not text:
            continue
        yield Document(
            page_content=text,
            metadata={"source": os.path.basename(path), "type": "pdf", "page": i + 1},
        )

def load_file(path: str) -> Iterable[Document]:
    ext = os.path.splitext(path)[1].lower()
    name = os.path.basename(path)

    if ext in PDF_EXTS:
        yield from load_pdf(path)
        return

    if ext in DOCX_EXTS:
        txt = load_docx(path)
    elif ext in TEXT_EXTS:
        txt = load_txt_like(path)
    else:
        return

    if txt:
        yield Document(
            page_content=txt,
            metadata={"source": name, "type": ext.lstrip(".")},
        )

def load_folder(corpus_dir: str) -> List[Document]:
    patterns = ["**/*.pdf", "**/*.docx", "**/*.txt", "**/*.md", "**/*.markdown", "**/*.html", "**/*.htm"]
    docs: List[Document] = []
    seen = set()
    for pat in patterns:
        for path in glob.glob(os.path.join(corpus_dir, pat), recursive=True):
            for d in load_file(path):
                key = (d.metadata.get("source"), d.metadata.get("page"), _hash_content(d.page_content[:1200]))
                if key in seen:
                    continue
                seen.add(key)
                docs.append(d)
    return docs

def load_urls(urls: Iterable[str]) -> List[Document]:
    out: List[Document] = []
    for url in urls:
        url = url.strip()
        if not url or url.startswith("#"):
            continue
        try:
            r = requests.get(url, timeout=20, headers={"User-Agent": "lifesync-lite/1.0"})
            r.raise_for_status()
            html = r.text

            if _HAS_READABILITY:
                article_html = ReadabilityDoc(html).summary(html_partial=True)
                soup = BeautifulSoup(article_html, "html.parser")
            else:
                soup = BeautifulSoup(html, "html.parser")
                for tag in soup(["script", "style", "noscript", "header", "footer", "nav", "aside"]):
                    tag.decompose()

            text = soup.get_text(separator=" ")
            text = _strip_boilerplate(text)
            if text:
                out.append(Document(page_content=text, metadata={"source": url, "type": "url"}))
        except Exception:
            # Skip bad URLs silently
            pass
    return out
